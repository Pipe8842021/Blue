# -*- coding: utf-8 -*-
"""Genera el documento Word del plan de trabajo en equipo de Blue Therapy."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colores de marca
TEAL = RGBColor(0x3A, 0xA8, 0x9E)
TEAL_DARK = RGBColor(0x1E, 0x40, 0x4D)
GOLD = RGBColor(0xA0, 0x7A, 0x28)
DARK = RGBColor(0x0D, 0x12, 0x14)
GRAY = RGBColor(0x4A, 0x65, 0x70)

doc = Document()

# Estilo base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x2A, 0x2E)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def set_cell_text(cell, text, bold=False, color=None, size=10, white=False, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = TEAL_DARK
    p.space_after = Pt(4)


def h2(text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = TEAL
    return p


def para(text, size=11, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = TEAL_DARK
    p.add_run(text)
    return p


# ---------- PORTADA ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('BLUE THERAPY')
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = TEAL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Plan de Trabajo en Equipo')
r.font.size = Pt(15)
r.font.color.rgb = GOLD
r.italic = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Mejora estética y funcional del sistema  ·  Ana, Jhonatan y Felipe')
r.font.size = Pt(11)
r.font.color.rgb = GRAY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Objetivo: pulir el diseño de todos los módulos, dejarlos más completos y '
              'agregar funciones útiles y sencillas, trabajando los tres a la vez sin '
              'que el trabajo de uno dependa del de otro.')
r.font.size = Pt(11)
r.italic = True
r.font.color.rgb = GRAY

doc.add_paragraph()

# ---------- 1. INTRODUCCION ----------
h1('1. ¿De qué va el proyecto?')
para('Blue Therapy es una aplicación web para administrar un centro de estética, spa y terapias. '
     'En una sola plataforma reúne tres cosas:')
bullet('el cliente entra, elige servicio y horario, y agenda solo (24/7).', 'Reservas en línea: ')
bullet('el centro confirma citas y gestiona clientes, servicios y galería.', 'Panel administrativo: ')
bullet('registra ingresos y egresos y muestra el balance del mes.', 'Control financiero: ')
para('Tecnología: PHP 8 + MySQL + Apache (XAMPP), con HTML, CSS y JavaScript propios (sin frameworks pesados). '
     'El sistema ya funciona de punta a punta; ahora toca hacerlo más bonito, más completo y con mejores funciones.',
     color=GRAY)

# ---------- 2. REGLAS ----------
h1('2. Reglas de oro para trabajar en paralelo')
para('Para que nadie dañe el trabajo del otro y avancemos al mismo tiempo:')
reglas = [
    ('Cada quien su rama de Git. ', 'Nadie trabaja en main. Cada rama sale de main: Ana → rama_ana, '
     'Jhonatan → rama_jhonatan, Felipe → rama_felipe (las tres ya están creadas en GitHub).'),
    ('Cada quien edita SOLO sus archivos. ', 'Cada uno tiene su lista exacta (ver punto 4).'),
    ('No tocar los archivos compartidos: ', 'config/db.php, includes/ (session, functions, auth, layout, footer), '
     'assets/css/admin.css y global.css. Esos se ajustan una sola vez, en acuerdo.'),
    ('Estilos nuevos → tu propio archivo CSS. ', 'No editar admin.css; cada quien agrega en su m-*.css.'),
    ('Funciones PHP nuevas → helper nombrado por lo que hace (no por la persona). ',
     'Ej. includes/h-agenda.php, h-catalogo.php, h-finanzas.php. Cada función con nombre descriptivo '
     'de su tarea, ej. calcularTotalCita(), formatearTelefonoWhatsapp(), exportarFinanzasCsv().'),
    ('Cambios en la base de datos → tu propia migración. ', 'database/migrations/NN_modulo.sql; nunca editar blue.sql directo.'),
    ('Commits pequeños y seguido. ', 'Antes de empezar cada día: git pull de tu rama.'),
]
for pre, txt in reglas:
    bullet(txt, pre)

# ---------- 3. COMO EMPEZAR ----------
h1('3. Cómo empezar a trabajar (Ana y Jhonatan)')
para('La base ya está lista: las ramas, la carpeta de migraciones, los CSS de cada módulo y el '
     '“gancho” para cargarlos. No tienen que configurar nada de eso, solo seguir estos pasos.',
     color=GRAY)
h2('Paso a paso (la primera vez)')
for t in [
    'Tener XAMPP con Apache y MySQL encendidos, y el proyecto en c:\\xampp\\htdocs\\Blue.',
    'Crear la base de datos: en phpMyAdmin importar database/blue.sql (crea blue_db). MySQL: root, sin contraseña.',
    'Traer el proyecto y cambiarse a su rama: git clone … / cd Blue / git checkout rama_ana (o rama_jhonatan).',
    'Probar que corre: http://localhost/Blue/login.php con admin@blue.com / admin123.',
    'Cada día: git pull al empezar; git add + commit + push al terminar.',
]:
    bullet(t)

h2('Lo PRIMERO que deben decirle a Claude Code')
para('Claude Code lee el CLAUDE.md automáticamente, pero necesita que le digan quién es cada uno '
     'para saber cuáles son “sus” archivos. Al abrir el proyecto, preséntense pegando este mensaje:')

p = doc.add_paragraph()
r = p.add_run('Para Ana:')
r.bold = True
r.font.color.rgb = TEAL_DARK
r.font.size = Pt(10)
q = doc.add_paragraph()
q.paragraph_format.left_indent = Inches(0.3)
r = q.add_run('“Hola. Soy Ana, del equipo de Blue Therapy. Trabajo en la rama rama_ana y me encargo de los '
              'módulos Clientes, Servicios y Galería. Antes de proponer cambios, lee el CLAUDE.md y '
              'docs/PLAN_EQUIPO.md. Edita solo mis archivos: admin/clients.php, admin/services.php, '
              'admin/gallery.php, mi CSS assets/css/m-catalogo.css y mi migración '
              'database/migrations/01_catalogo.sql. No toques archivos compartidos ni de otros compañeros.”')
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GRAY

p = doc.add_paragraph()
r = p.add_run('Para Jhonatan:')
r.bold = True
r.font.color.rgb = TEAL_DARK
r.font.size = Pt(10)
q = doc.add_paragraph()
q.paragraph_format.left_indent = Inches(0.3)
r = q.add_run('“Hola. Soy Jhonatan, del equipo de Blue Therapy. Trabajo en la rama rama_jhonatan y me encargo de '
              'los módulos Reserva en línea, Finanzas y Login/Configuración. Antes de proponer cambios, lee el '
              'CLAUDE.md y docs/PLAN_EQUIPO.md. Edita solo mis archivos: booking.php, assets/js/booking.js, '
              'assets/css/booking.css, admin/finances.php, login.php, admin/settings.php, mi CSS '
              'assets/css/m-finanzas.css y mi migración database/migrations/02_finanzas.sql. No toques archivos '
              'compartidos ni de otros compañeros.”')
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = GRAY

# ---------- 4. REPARTO ----------
doc.add_page_break()
h1('4. Reparto de tareas (equitativo)')
para('Cada persona tiene 3 áreas: una pesada, una media y una ligera, para que el esfuerzo sea parejo.', color=GRAY)

people = [
    ('FELIPE — Agenda y Tablero', [
        ('Citas (pesada)', 'Pantalla donde el centro ve, filtra y administra todas las citas y su estado.', 'Vista de calendario semanal, botón “recordar por WhatsApp”, imprimir cita.'),
        ('Dashboard (media)', 'Pantalla de inicio del panel con los indicadores clave del negocio.', 'Mini-gráfico de citas por día (Chart.js), próximas citas, Top 5 servicios.'),
        ('Panel Staff (ligera)', 'Vista del profesional con su agenda y las citas que le asignaron.', 'Agenda del día del profesional con sus citas asignadas.'),
    ], 'm-agenda.css  ·  03_agenda.sql'),
    ('ANA — Catálogo y Clientes', [
        ('Clientes (pesada)', 'Listado y ficha de cada cliente con su historial e información de contacto.', 'Ficha detallada con historial (timeline), total gastado, etiquetas (VIP/Nuevo/Frecuente).'),
        ('Servicios (media)', 'Catálogo de servicios y categorías que ofrece el centro.', 'Imagen por servicio, marcar “destacado”, filtro por categoría.'),
        ('Galería (ligera)', 'Imágenes del centro que se muestran en el sitio público.', 'Lightbox al hacer clic, agrupar por categoría.'),
    ], 'm-catalogo.css  ·  01_catalogo.sql'),
    ('JHONATAN — Finanzas, Reservas y Acceso', [
        ('Reserva en línea (pesada)', 'Asistente público donde el cliente agenda su cita paso a paso.', 'Validación en vivo, selección visual de servicios, barra de progreso mejorada.'),
        ('Finanzas (media)', 'Registro de ingresos y egresos con el resumen y balance del mes.', 'Gráfico ingresos vs egresos (Chart.js), filtro por rango de fechas, exportar a CSV.'),
        ('Login + Configuración (ligera)', 'Acceso al sistema y ajustes de perfil y del equipo.', 'Mostrar/ocultar contraseña, avatar de usuario, fortaleza de contraseña.'),
    ], 'm-finanzas.css  ·  02_finanzas.sql'),
]

for name, rows, files in people:
    h2(name)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for c, txt in zip(hdr, ['Área', 'De qué trata', 'Funciones nuevas (sencillas)']):
        set_cell_text(c, txt, bold=True, white=True, size=10)
        shade(c, '3AA89E')
    for area, desc, func in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], area, bold=True, color=TEAL_DARK, size=10)
        set_cell_text(cells[1], desc, size=9, color=GRAY)
        set_cell_text(cells[2], func, size=10)
    p = doc.add_paragraph()
    r = p.add_run('Tu CSS y migración: ' + files)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GOLD

h2('Reglas para nombrar tus archivos nuevos')
para('Las páginas que ya existen se editan tal cual. Si creas un archivo nuevo, nómbralo así:', size=10)
for t in [
    'Describe lo que hace, no quién lo hizo (nada de pagina_ana.php).',
    'Minúsculas, sin tildes ni espacios; separa palabras con guion bajo (reporte_servicios.php).',
    'Mantén el estilo de los que ya existen (appointments.php, clients.php, finances.php…).',
    'Helpers de PHP: includes/h-<tarea>.php, y cada función con nombre descriptivo (calcularTotalCita(), exportarFinanzasCsv()).',
    'Imágenes: en tu subcarpeta propia dentro de assets/img/ (ej. assets/img/servicios/).',
    'Tu CSS y tu migración ya están creados: solo agregas dentro.',
]:
    bullet(t)

para('Nota: la landing pública y el estilo base ya quedaron listos en la preparación común, para mantener '
     'una misma identidad visual en todo el sitio.', size=10, color=GRAY)

# ---------- 5. BUENAS PRACTICAS / DISENO ----------
doc.add_page_break()
h1('5. Diseño y buenas prácticas (para que todo combine)')
h2('Colores e identidad')
bullet('Turquesa de marca (#5BC4D0 / teal #5BC4B8), dorado premium con moderación, oscuros del logo.', 'Paleta: ')
bullet('Títulos en Playfair Display; texto en DM Sans.', 'Tipografía: ')
bullet('Usar SIEMPRE las variables de color existentes; no inventar colores nuevos ni hex sueltos.', 'Regla: ')
h2('Estética')
for t in [
    'Limpio, profesional y premium (estilo spa), con espaciado generoso.',
    'Responsive siempre (probar en celular y computador).',
    'Mismos botones, tarjetas y tipografía en los 3 módulos.',
    'Microinteracciones sutiles (hover, transiciones suaves), sin exageraciones.',
    'Librerías externas solo por CDN y solo si son simples (ej. Chart.js).',
]:
    bullet(t)
h2('Código')
for t in [
    'Consultas con PDO preparado (nunca concatenar SQL) y escapar la salida con e().',
    'Formularios con token CSRF; mensajes con setFlash() y redirigir tras un POST.',
    'Verificar con php -l y probar en el navegador antes de dar por terminado.',
]:
    bullet(t)

# ---------- 6. TERMINADO ----------
h1('6. “Terminado” significa')
for t in [
    'Se ve bien en computador y celular (responsive).',
    'Sin errores de PHP.',
    'Formularios validan y muestran mensajes claros.',
    'Probado en el navegador.',
    'Coherente con el diseño común.',
    'Si cambiaste la base de datos, está en tu archivo de migración.',
]:
    bullet(t)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run('Blue Therapy · Documento de equipo · junio 2026')
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = GRAY

import sys
out = sys.argv[1] if len(sys.argv) > 1 else r'c:\xampp\htdocs\Blue\docs\Blue_Therapy_Plan_Equipo.docx'
doc.save(out)
print('OK ->', out)
